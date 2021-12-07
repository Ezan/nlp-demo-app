from functools import reduce
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from nltk.tokenize import word_tokenize, sent_tokenize

class DocReader:
    def __init__(self, source):
        self.doc = Document(source)

    def _highlighter(self, phrase):
        for para in self.doc.paragraphs :
            start = para.text.find(phrase)
            if start > -1 :
                pre = para.text[:start]
                post = para.text[start+len(phrase):]
                para.text = pre
                para.add_run(phrase)
                para.runs[1].font.highlight_color = WD_COLOR_INDEX.YELLOW
                para.add_run(post)
                

    def highlight(self, savefile, result):
        phrase_list = list(filter(lambda p: p != '', reduce(lambda x, y: x+list(map(lambda z: z['answer'], y['result'])), result, [])))
        for item in phrase_list:
            self._highlighter(item)
        self.doc.save(savefile)
       
    def extract(self):
        contexts = [para.text for para in self.doc.paragraphs]
        return contexts
